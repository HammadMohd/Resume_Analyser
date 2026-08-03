"""Tests for parser module."""

import os
import tempfile

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.parser.docx_parser import DOCXParser
from backend.parser.pdf_parser import PDFParser
from backend.parser.resume_parser import ResumeParser


@pytest.fixture
def pdf_file():
    """Create a temporary PDF file for testing."""
    # Create a minimal PDF
    content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << >> >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (Test Resume) Tj ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
trailer
<< /Size 5 /Root 1 0 R >>
startxref
360
%%EOF"""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(content)
        return f.name


@pytest.fixture
def docx_file():
    """Create a temporary DOCX file for testing."""
    doc = Document()
    doc.add_heading("John Doe", level=1)
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Experience: 5 years in Python development")
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path


class TestPDFParser:
    """Tests for PDF parser."""

    def test_parse_returns_result(self, pdf_file):
        """PDF parser should return ParsedResume."""
        parser = PDFParser()
        result = parser.parse(pdf_file, "test.pdf")
        assert result.filename == "test.pdf"
        assert result.parser_used == "pdfplumber"
        assert result.total_pages >= 1
        os.unlink(pdf_file)


class TestDOCXParser:
    """Tests for DOCX parser."""

    def test_parse_returns_result(self, docx_file):
        """DOCX parser should return ParsedResume."""
        parser = DOCXParser()
        result = parser.parse(docx_file, "test.docx")
        assert result.filename == "test.docx"
        assert result.parser_used == "python-docx"
        assert "John Doe" in result.full_text
        os.unlink(docx_file)

    def test_extracts_headers(self, docx_file):
        """DOCX parser should detect headers."""
        parser = DOCXParser()
        result = parser.parse(docx_file, "test.docx")
        headers = [b for b in result.pages[0].text_blocks if b.is_header]
        assert len(headers) > 0
        os.unlink(docx_file)


class TestResumeParser:
    """Tests for resume parser coordinator."""

    def test_parse_pdf(self, pdf_file):
        """Resume parser should handle PDF files."""
        parser = ResumeParser()
        result = parser.parse(pdf_file, "test.pdf")
        assert result.filename == "test.pdf"
        os.unlink(pdf_file)

    def test_parse_docx(self, docx_file):
        """Resume parser should handle DOCX files."""
        parser = ResumeParser()
        result = parser.parse(docx_file, "test.docx")
        assert result.filename == "test.docx"
        os.unlink(docx_file)

    def test_unsupported_type(self):
        """Resume parser should reject unsupported types."""
        parser = ResumeParser()
        with pytest.raises(Exception):
            parser.parse("test.txt", "test.txt")


class TestParseEndpoint:
    """Tests for parse endpoint."""

    def test_parse_pdf_endpoint(self, client: TestClient, pdf_file):
        """Parse endpoint should accept PDF uploads."""
        with open(pdf_file, "rb") as f:
            files = {"file": ("test.pdf", f, "application/pdf")}
            response = client.post("/api/v1/resumes/parse", files=files)
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "parsed" in data["data"]
        os.unlink(pdf_file)

    def test_parse_docx_endpoint(self, client: TestClient, docx_file):
        """Parse endpoint should accept DOCX uploads."""
        with open(docx_file, "rb") as f:
            files = {"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = client.post("/api/v1/resumes/parse", files=files)
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "parsed" in data["data"]
        os.unlink(docx_file)
