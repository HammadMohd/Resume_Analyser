"""ATS-Compliant DOCX Resume Exporter using python-docx.

Generates clean single-column DOCX resume files.
"""

import io

import docx
from docx.shared import Inches, Pt

from backend.schemas.resume import NormalizedResume
from backend.utils.logging import get_logger

logger = get_logger(__name__)


class DOCXResumeExporter:
    """Exports structured resume into an ATS-compliant DOCX format."""

    def export_docx(self, resume: NormalizedResume) -> bytes:
        """Generate DOCX binary content from normalized resume."""
        doc = docx.Document()

        # Set 0.5 inch margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Header Name
        heading = doc.add_heading(level=1)
        run = heading.add_run(resume.contact.name or "Candidate Resume")
        run.font.size = Pt(18)
        run.font.bold = True

        # Contact Info
        contact_p = doc.add_paragraph()
        contact_str = f"{resume.contact.email or ''} | {resume.contact.phone or ''} | {resume.contact.location or ''}"
        contact_p.add_run(contact_str)

        # Experience
        if resume.experience:
            exp_h = doc.add_heading("WORK EXPERIENCE", level=2)
            exp_h.runs[0].font.size = Pt(12)
            for exp in resume.experience:
                p = doc.add_paragraph()
                r1 = p.add_run(f"{exp.title} - {exp.company} ")
                r1.bold = True
                date_str = f"{exp.start_date} - {exp.end_date}".strip(" -") if (exp.start_date or exp.end_date) else ""
                if date_str:
                    r2 = p.add_run(f"({date_str})")
                    r2.italic = True

                for b in exp.bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.add_run(b)

        # Education
        if resume.education:
            edu_h = doc.add_heading("EDUCATION", level=2)
            edu_h.runs[0].font.size = Pt(12)
            for edu in resume.education:
                p = doc.add_paragraph()
                edu_date = f"{edu.start_date} - {edu.end_date}".strip(" -") if (edu.start_date or edu.end_date) else ""
                p.add_run(f"{edu.degree} - {edu.institution} ({edu_date})" if edu_date else f"{edu.degree} - {edu.institution}")

        # Skills
        if resume.skills:
            skills_h = doc.add_heading("SKILLS", level=2)
            skills_h.runs[0].font.size = Pt(12)
            sp = doc.add_paragraph()
            skills_str = ", ".join([s for cat in resume.skills for s in cat.skills])
            sp.add_run(skills_str)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes
